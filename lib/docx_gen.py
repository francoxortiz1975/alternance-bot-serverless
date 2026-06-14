"""Génération de la lettre de motivation DOCX à partir du template + infos Gemini."""

import copy
from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .gemini_client import raccourcir_para_tech

TEMPLATE_LETTRE = Path(__file__).parent.parent / "templates" / "MODEL_Lettre_Motivation_ORTIZ_IngData_.docx"

PORTFOLIO_URL = "https://www.francoxortiz.vercel.app"
PORTFOLIO_TEXTE = "francoxortiz.dev"


def date_fr():
    mois = ["janvier", "février", "mars", "avril", "mai", "juin",
            "juillet", "août", "septembre", "octobre", "novembre", "décembre"]
    d = date.today()
    return f"{d.day} {mois[d.month-1]} {d.year}"


def remplacer_paragraphe_complet(para, remplacements):
    texte = para.text
    modifie = False
    for ancien, nouveau in remplacements.items():
        if ancien in texte:
            texte = texte.replace(ancien, nouveau)
            modifie = True
    if modifie and para.runs:
        para.runs[0].text = texte
        for run in para.runs[1:]:
            run.text = ""


def remplacer_paragraphe_technique(doc, nouveau_texte):
    for para in doc.paragraphs:
        if para.text.strip().startswith("Sur le plan"):
            if para.runs:
                para.runs[0].text = nouveau_texte
                for run in para.runs[1:]:
                    run.text = ""
            return True
    return False


def compter_mots_lettre(doc):
    """Compte les mots de tous les paragraphes non vides."""
    return sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())


def ajouter_hyperlien_portfolio(doc):
    """Transforme le texte affiché 'francoxortiz.dev' en hyperlien vers francoxortiz.vercel.app."""
    for para in doc.paragraphs:
        if PORTFOLIO_TEXTE not in para.text:
            continue

        r_id = doc.part.relate_to(
            PORTFOLIO_URL,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        for run in para.runs:
            if PORTFOLIO_TEXTE not in run.text:
                continue
            new_run = copy.deepcopy(run._r)
            rPr = new_run.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                new_run.insert(0, rPr)
            rStyle = OxmlElement("w:rStyle")
            rStyle.set(qn("w:val"), "Hyperlink")
            rPr.insert(0, rStyle)
            hyperlink.append(new_run)
            para._p.replace(run._r, hyperlink)
            break
        return True
    return False


def generer_lettre_bytes(infos):
    """Génère la lettre de motivation DOCX et retourne les bytes du fichier."""
    doc = Document(TEMPLATE_LETTRE)

    remplacements = {
        "XXX_ENTREPRISENAME":     infos["entreprise_nom"],
        "XXX_DEPARTEMENT":        infos["equipe_departement"],
        "XXX_LOCALISATION":       infos["localisation"],
        "Paris, le XXX_DAYOFTHEMONTH juin 2026": f"Paris, le {date_fr()}",
        "XXX_OFFERTITLE":         infos["titre_objet"],
        "XXX_TEXTRELATEDTOTHEOFFER": infos["phrase_motivation"],
        "XXX_DOMAINOFTHEOFFER":   infos["domaine_technique"],
        "XXX_FIELDOFTHEOFFER":    infos["secteur"],
    }

    for para in doc.paragraphs:
        remplacer_paragraphe_complet(para, remplacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    remplacer_paragraphe_complet(para, remplacements)

    para_tech = infos.get("paragraphe_tech_adapte", "")
    if para_tech:
        remplacer_paragraphe_technique(doc, para_tech)

    # Check word count — raccourcit directement si > 400 mots
    mots = compter_mots_lettre(doc)
    if mots > 400:
        para_actuel = infos.get("paragraphe_tech_adapte", "")
        nouveau_para = raccourcir_para_tech(infos, max(300, len(para_actuel) - 300))
        remplacer_paragraphe_technique(doc, nouveau_para)

    # Hyperlien portfolio cliquable
    ajouter_hyperlien_portfolio(doc)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
